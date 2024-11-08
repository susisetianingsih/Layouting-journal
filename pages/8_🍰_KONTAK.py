import streamlit as st
from docx import Document
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime
import pythoncom
import win32com.client as win32

st.set_page_config(
    page_title="KONTAK",
    page_icon="🍰"
)

def extract_info_from_docx(file):
    doc = Document(file)
    journal_info = {}

    paragraphs = doc.paragraphs
    if len(paragraphs) > 0:
        journal_info["title"] = paragraphs[0].text
    if len(paragraphs) > 2:
        journal_info["author"] = paragraphs[2].text
    if len(paragraphs) > 4:
        journal_info["afiliasi"] = paragraphs[4].text
    
    return journal_info

def fill_template(template_file, journal_info, loa_id, current_date, loa_name):
    doc = Document(template_file)
    
    # Replace placeholders with actual data
    for paragraph in doc.paragraphs:
        ROMAWI_BULAN = "VI"
        TAHUN = "2024"
        VOL_NO = "Volume 1 Nomor 1 Juni 2024"
        
        formatted_authors = format_authors(journal_info["author"], journal_info["afiliasi"])
        if '{judul}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{judul}', journal_info["title"])
        if '{romawi_bulan}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{romawi_bulan}', ROMAWI_BULAN)
        if '{tahun}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{tahun}', TAHUN)
        if '{vol_no}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{vol_no}', VOL_NO)
        if '{penulis}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{penulis}', formatted_authors)
        if '{loa_id}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{loa_id}', loa_id)
        if '{tanggal}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{tanggal}', current_date)
            
        apply_style_to_paragraph(paragraph)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{judul}' in cell.text:
                        cell.text = cell.text.replace('{judul}', journal_info["title"])
                        for paragraph in cell.paragraphs:
                            apply_style_to_paragraph(paragraph, bold=True)
                    if '{penulis}' in cell.text:
                        cell.text = cell.text.replace('{penulis}', formatted_authors)
                    
                    # Apply font and size to each cell
                    for paragraph in cell.paragraphs:
                        apply_style_to_paragraph(paragraph)
    
    # Save the filled document
    output_path_word = os.path.join("D:\\PROJECT\\KONTAK\\LoA", f"{loa_name}.docx")
    output_path_pdf = os.path.join("D:\\PROJECT\\KONTAK\\LoA", f"{loa_name}.pdf")
        
    doc.save(output_path_word)
    # Convert DOCX to PDF using win32com
    pythoncom.CoInitialize()  # Ensure COM initialization
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(output_path_word)
        doc.SaveAs(output_path_pdf, FileFormat=17)  # 17 is the format code for PDF
        doc.Close()
        word.Quit()
    except Exception as e:
        return f"Error in PDF conversion: {e}"
    finally:
        pythoncom.CoUninitialize()
    return "File Word dan PDF berhasil dibuat 🤩"

def format_authors(authors, affiliation):
    # Extract university name from the affiliation string
    university_name = re.search(r'\b(Sekolah|Institut|Politeknik|Universitas|University|Poltekkes)\b.*?, Indonesia', affiliation).group(0)[:-11]
    authors_list = [re.sub(r'[\*\d]+$', '', author.strip()) for author in authors.split(',')]
    
    # Create formatted author strings with numbering
    formatted_authors = []
    for idx, author in enumerate(authors_list, start=1):
        formatted_authors.append(f"{idx}. {author}, {university_name}")
    
    return '\n'.join(formatted_authors)

def extract_loa_name(file_name):
    # Gunakan regex untuk mengekstrak ID dan nama
    match = re.search(r'ID(\d+)\s+([^\d_]+)', file_name)
    if match:
        loa_id = match.group(1)
        loa_name = match.group(2).strip()
        return loa_id, f"LOA KONTAK_{loa_id} {loa_name}"
    return "Nama LoA tidak ditemukan"

def apply_style_to_paragraph(paragraph, bold=False):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
        if bold:
            run.bold = True
    # Set single line spacing and 0 pt spacing after the paragraph
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # Single line spacing
    paragraph_format.space_after = Pt(0)
    

def get_current_date():
    now = datetime.now()
    return f"{int(now.strftime('%d'))} {now.strftime('%B %Y')}"

# Membuat antarmuka Streamlit
st.title("Ekstraksi Data KONTAK")
st.write("Unggah file jurnal Word (DOCX).")

# Upload file
uploaded_file = st.file_uploader("Pilih file Word", type=["docx"])

if uploaded_file is not None:
    # Ambil nama file untuk mengekstrak LoA
    file_name = uploaded_file.name
    loa_id, loa_name = extract_loa_name(file_name)
    TEMPLATE_PATH = "templates/LOA KONTAK_.docx"

    # Ekstrak informasi dari file yang diunggah
    journal_info = extract_info_from_docx(uploaded_file)
    
    # Dapatkan tanggal saat ini
    current_date = get_current_date()

    # Tampilkan hasil ekstraksi
    st.subheader("Hasil Ekstraksi:")
    st.write("**Judul:**", journal_info["title"])
    st.write("**Penulis:**", journal_info["author"])
    st.write("**Afiliasi:**", journal_info["afiliasi"])
    st.write("**Penamaan LoA:**", loa_name)
    st.write("**LoA ID:**", loa_id)

    # Isi template dengan data yang diekstrak
    filled_file = fill_template(TEMPLATE_PATH, journal_info, loa_id, current_date, loa_name)
    st.success(filled_file)
    
#     # Provide download link for the filled document
#     with open(filled_word_file, "rb") as file:
#         st.download_button(
#             label="Unduh Dokumen LoA",
#             data=file,
#             file_name=filled_word_file,
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

#     # Clean up the generated file if necessary
#     os.remove(filled_word_file)
