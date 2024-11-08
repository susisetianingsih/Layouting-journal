import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from io import BytesIO
import re
from docx.shared import RGBColor

# Check if a paragraph contains an image by looking for 'inline' elements
def contains_image(paragraph):
    for run in paragraph.runs:
        if run._element.findall('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
            return True
    return False

# Fungsi untuk format paragraf sesuai aturan yang diberikan
def format_paragraph(paragraph, font_name="Cambria", font_size=11, italic=False, bold=False, alignment='justified', 
                     left_indent=0, right_indent=0, spacing_before=0, spacing_after=0, 
                     line_spacing=1, special_indent="none"):
    
    # Set font
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Set alignment
    if alignment == 'justified':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    elif alignment == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif alignment == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif alignment == 'left':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Set indentasi
    paragraph.paragraph_format.left_indent = Pt(left_indent)
    paragraph.paragraph_format.right_indent = Pt(right_indent)
    paragraph.paragraph_format.space_before = Pt(spacing_before)
    paragraph.paragraph_format.space_after = Pt(spacing_after)

    paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    
    # Line spacing
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Set special indentasi untuk "first line"
    if special_indent == "first_line":
        paragraph.paragraph_format.first_line_indent = Pt(36)  # Set first line indent to 1.27 cm
    else:
        paragraph.paragraph_format.first_line_indent = None  # No first line indent if not specified

def reformat_document(doc):
    to_delete = []  # List untuk menandai paragraf yang akan dihapus
    abstrak_section = False  # Menandai apakah sedang dalam bagian Abstrak hingga Keywords
    afiliation_section = False  # Menandai apakah sedang dalam bagian Abstrak hingga Keywords
    # daftar_pustaka_section = False  # Menandai apakah telah mencapai bagian "DAFTAR PUSTAKA"
    
    # Hapus paragraf kosong atau dengan "nomor handphone" dari dokumen
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().lower()
        
        if contains_image(paragraph):
            continue
        if not text:
            to_delete.append(paragraph)
        elif "nomor handphone" in text:
            to_delete.append(paragraph)

    # Menghapus paragraf yang ditandai
    for paragraph in to_delete:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    
    # Menambahkan paragraf kosong setelah "abstrak", "abstract", "gambar {angka}", "tabel {angka}", dan paragraf kedua
    pattern = re.compile(r'\b(kata kunci|keywords|abstrak|abstract|gambar \d. |tabel \d. )\b', re.IGNORECASE)
    paragraphs_to_insert_before = []
    
    # Tandai paragraf yang cocok dan paragraf kedua
    for i, paragraph in enumerate(doc.paragraphs):
        if contains_image(paragraph):
            continue  # Skip formatting if paragraph contains an image
        
        format_paragraph(paragraph, spacing_before=0, spacing_after=0, special_indent="first_line")
        
        if i == 2:  # Paragraf kedua (indeks 1)
            paragraphs_to_insert_before.append(i)
        if pattern.search(paragraph.text):
            format_paragraph(paragraph)
            if 'gambar' in text:
                i+=1
            paragraphs_to_insert_before.append(i)

    # Sisipkan paragraf kosong setelah paragraf yang sesuai
    for index in reversed(paragraphs_to_insert_before):
        empty_paragraph = doc.paragraphs[index].insert_paragraph_before("")
        format_paragraph(empty_paragraph, font_size=10, special_indent="none")

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().lower()  
        
        if contains_image(paragraph):
            continue  # Skip formatting if paragraph contains an image

        if i == 0:
            paragraph.text = paragraph.text.title()  # Kapitalisasi setiap kata
            format_paragraph(paragraph, font_size=14, special_indent="none", spacing_after=18, alignment='center')
        elif i == 1:
            paragraph.text = paragraph.text.title()
            format_paragraph(paragraph, font_size=12, bold=True, alignment='center', special_indent="none")
            afiliation_section = True
            
        # Paragraf dengan "Abstrak" atau "Abstract"
        elif 'abstrak' in text or 'abstract' in text:
            afiliation_section = False
            abstrak_section = True  # Mulai section Abstrak
            format_paragraph(paragraph, font_size=10, italic=True, bold=True, alignment='center', special_indent="none")
        
        elif afiliation_section:
            format_paragraph(paragraph, font_size=10, alignment='center', special_indent="none")
        
        elif 'kata kunci' in text or 'keywords' in text:
            format_paragraph(paragraph, font_size=10, italic=True, alignment='left', special_indent="none")
            if 'keywords:' in text:
                abstrak_section = False
            
        #Jika dalam section Abstrak hingga Keywords, gunakan font size 10pt
        elif abstrak_section:
            format_paragraph(paragraph, italic=True, font_size=10, special_indent="first_line")
    
    # Kompilasi pola header dan sub-header
    pattern_header = re.compile(r'\b(\d. PENDAHULUAN|\d. METODE|\d. HASIL DAN PEMBAHASAN|\d. KESIMPULAN|UCAPAN TERIMA KASIH|DAFTAR PUSTAKA)\b', re.IGNORECASE)
    pattern_sub_header = re.compile(r'\b(\d. \d.|\d. \d. \d.)\b', re.IGNORECASE)

    # Loop melalui semua paragraf di dokumen
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().lower()

        # Format header
        if pattern_header.search(text):
            format_paragraph(paragraph, bold=True, spacing_before=24, spacing_after=6)
            # if "daftar pustaka" in text:
                # daftar_pustaka_section = True  # Mulai section Daftar Pustaka
        
        # Format sub-header
        elif pattern_sub_header.search(text):
            format_paragraph(paragraph, bold=True, spacing_before=12, spacing_after=3)
        
        # Terapkan indentasi hanging jika berada dalam section Daftar Pustaka
        # elif daftar_pustaka_section:
        #     format_paragraph(paragraph, special_indent="hanging", spacing_after=3)
        #     paragraph.paragraph_format.left_indent = Cm(0.84)  # Set hanging indent to 0.84 cm


    return doc

# Streamlit UI
st.title("Word Document Type JAMSI")
uploaded_file = st.file_uploader("Upload a Word file (.docx)", type="docx")

if uploaded_file is not None:
    doc = Document(uploaded_file)
    formatted_doc = reformat_document(doc)

    # Save the formatted document to an in-memory buffer
    buffer = BytesIO()
    formatted_doc.save(buffer)
    buffer.seek(0)
    file_name = uploaded_file.name.replace(".docx", "")

    # Provide download link
    st.download_button(label="Download Formatted Document", data=buffer, file_name=f"{file_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
